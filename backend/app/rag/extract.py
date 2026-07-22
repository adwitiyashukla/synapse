"""Text extraction from uploaded files (PDF, DOCX, TXT, MD)."""

import io

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


class ExtractionError(Exception):
    pass


def extract_text(filename: str, data: bytes) -> str:
    """Extract plain text from a file based on its extension."""
    lowered = filename.lower()
    if lowered.endswith(".pdf"):
        return _extract_pdf(data)
    if lowered.endswith(".docx"):
        return _extract_docx(data)
    if lowered.endswith((".txt", ".md")):
        return _extract_plain(data)
    raise ExtractionError(
        "Unsupported file type. Allowed: " + ", ".join(sorted(ALLOWED_EXTENSIONS))
    )


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
    except Exception as exc:
        raise ExtractionError(f"Could not read PDF: {exc}") from exc
    text = "\n\n".join(pages).strip()
    if not text:
        raise ExtractionError(
            "No extractable text found in the PDF. It may be a scanned image."
        )
    return text


def _extract_docx(data: bytes) -> str:
    import docx

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise ExtractionError(f"Could not read DOCX: {exc}") from exc
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    text = "\n".join(part for part in parts if part.strip()).strip()
    if not text:
        raise ExtractionError("No text found in the DOCX file.")
    return text


def _extract_plain(data: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            text = data.decode(encoding).strip()
            if text:
                return text
        except UnicodeDecodeError:
            continue
    raise ExtractionError("Could not decode the text file.")
